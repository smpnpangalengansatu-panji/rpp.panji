import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- KONFIGURASI HALAMAN ---
st.set_page_config(
    page_title="AI Modul Ajar Expert 2026", 
    page_icon="🎓",
    layout="wide"
)

# --- CUSTOM CSS UNTUK TAMPILAN PROFESIONAL ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    .main {
        background-color: #f4f7f9;
    }
    div[data-testid="stForm"] {
        background-color: #ffffff;
        padding: 30px;
        border-radius: 15px;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
        border: none;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3.5em;
        background-color: #007bff;
        color: white;
        font-weight: 600;
        border: none;
        transition: 0.3s;
    }
    .stButton>button:hover {
        background-color: #0056b3;
        box-shadow: 0 4px 12px rgba(0,123,255,0.3);
    }
    h1, h2, h3 {
        color: #1e293b;
    }
    </style>
    """, unsafe_allow_html=True)

# --- FUNGSI PARSER DOCX ---
def create_formatted_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    lines = text.split('\n')
    is_table = False
    table_data = []

    for line in lines:
        clean_line = line.strip()
        if '|' in clean_line:
            if '---' in clean_line: continue
            cells = [c.strip() for c in clean_line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
                is_table = True
            continue
        else:
            if is_table and table_data:
                table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                table.style = 'Table Grid'
                for i, row in enumerate(table_data):
                    for j, cell_text in enumerate(row):
                        table.rows[i].cells[j].text = cell_text.replace('**', '')
                table_data = []
                is_table = False
                doc.add_paragraph()
        if not clean_line: continue
        if clean_line.startswith('#'):
            level = clean_line.count('#')
            doc.add_heading(clean_line.replace('#', '').strip(), level=min(level, 3))
        elif clean_line.startswith(('* ', '- ')):
            doc.add_paragraph(clean_line[2:].replace('**', ''), style='List Bullet')
        elif clean_line[0:1].isdigit() and clean_line[1:2] == '.':
            doc.add_paragraph(clean_line[2:].strip().replace('**', ''), style='List Number')
        elif clean_line.startswith('**') and clean_line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(clean_line.replace('**', ''))
            run.bold = True
        else:
            doc.add_paragraph(clean_line.replace('**', ''))
    return doc

# --- SIDEBAR ---
with st.sidebar:
    st.markdown("### 🎓 Admin Kurikulum AI")
    st.image("https://cdn-icons-png.flaticon.com/512/3976/3976625.png", width=100)
    st.info("Aplikasi Generate Modul Ajar Kurikulum Merdeka. by: 'IN2026'")
    st.divider()
    st.caption("v2.6 - Auto-Model Resilience")

# --- SESSION STATE ---
if 'page' not in st.session_state:
    st.session_state.page = 1
if 'data' not in st.session_state:
    st.session_state.data = {}

def go_to_page(page_number):
    st.session_state.page = page_number
    st.rerun()

# --- HALAMAN 1: INPUT DATA ---
if st.session_state.page == 1:
    st.title("📝 Penyusunan Modul Ajar")
    st.subheader("Bagian 1: Identitas")
    
    # Ambil data lama jika ada, jika tidak ada gunakan string kosong atau default
    old_data = st.session_state.get('data', {})

    with st.form("form_halaman_1"):
        st.markdown("#### 🔑 Konfigurasi API")
        api_key = st.text_input("Masukkan API Key Gemini:", 
                                type="password", 
                                value=old_data.get('api_key', "")) # Data tetap ada
        
        st.markdown("#### 👤 Identitas Guru & Kelas")
        col1, col2 = st.columns(2)
        with col1:
            nama = st.text_input("Nama Guru", 
                                 value=old_data.get('nama', ""), 
                                 placeholder="Contoh: Iman Nuriman, ST.")
            unit = st.text_input("Unit Kerja", 
                                 value=old_data.get('unit', ""), 
                                 placeholder="Contoh: SMP Negeri 1 Pangalengan")
            mapel = st.text_input("Mata Pelajaran", 
                                  value=old_data.get('mapel', ""))
            
            # Untuk selectbox, cari indeksnya agar tidak kembali ke pilihan pertama
            fase_options = ["A", "B", "C", "D", "E", "F"]
            try:
                fase_idx = fase_options.index(old_data.get('fase'))
            except:
                fase_idx = 0
            fase = st.selectbox("Fase", fase_options, index=fase_idx)
        
        with col2:
            kelas = st.text_input("Kelas", value=old_data.get('kelas', ""))
            
            sem_options = ["1 (Ganjil)", "2 (Genap)"]
            try:
                sem_idx = sem_options.index(old_data.get('semester'))
            except:
                sem_idx = 0
            semester = st.selectbox("Semester", sem_options, index=sem_idx)
            
            jp = st.text_input("Alokasi Waktu", 
                               value=old_data.get('jp', ""), 
                               placeholder="Contoh: 2 x 40 Menit")
            topik = st.text_input("Topik Pembelajaran", 
                                  value=old_data.get('topik', ""))

        st.markdown("#### 🎯 Dimensi Profil Lulusan (DPL)")
        # Untuk multiselect
        dpl_list = ["Keimanan", "Wargaan", "Penalaran Kritis", "Kreativitas", "Kolaborasi", "Kemandirian", "Kesehatan", "Komunikasi"]
        dimensi_dpl = st.multiselect(
            "Pilih Dimensi Capaian:",
            dpl_list,
            default=old_data.get('dimensi_dpl', ["Penalaran Kritis", "Kreativitas"])
        )

        st.markdown("#### ⚙️ Metode Pembelajaran")
        model_options = [
            "Pembelajaran Berbasis Masalah (PBL)", 
            "Pembelajaran Berbasis Proyek (PjBL)", 
            "Inquiry Learning", 
            "Cooperative Learning",
            "Discovery Learning",
            "Pembelajaran Berdiferensiasi",
            "Pilihkan secara otomatis oleh AI"
        ]
        try:
            model_idx = model_options.index(old_data.get('model'))
        except:
            model_idx = 0
        model_belajar = st.selectbox("Model Pembelajaran", model_options, index=model_idx)
        
        pertemuan = st.number_input("Jumlah Pertemuan", min_value=1, value=old_data.get('pertemuan', 1))

        st.markdown("#### 📝 Kondisi Tambahan (Opsional)")
        kondisi_khusus = st.text_area(
            "Tambahkan instruksi khusus:",
            value=old_data.get('kondisi_khusus', ""),
            placeholder="Contoh: Hubungkan materi dengan budaya lokal..."
        )

        submit_1 = st.form_submit_button("Lanjut ke Konfirmasi →")
        
               
        if submit_1:
            if not api_key or not nama or not topik:
                st.error("Mohon lengkapi data wajib (API Key, Nama, dan Topik)!")
            else:
                st.session_state.data = {
                    'api_key': api_key, 'nama': nama, 'unit': unit, 'mapel': mapel,
                    'fase': fase, 'kelas': kelas, 'semester': semester,
                    'jp': jp, 'pertemuan': pertemuan, 'topik': topik, 'model': model_belajar,
                    'dimensi_dpl': dimensi_dpl,
                    'kondisi_khusus': kondisi_khusus # Menyimpan kondisi tambahan
                }
                go_to_page(2)

# --- HALAMAN 2: KONFIRMASI ---
elif st.session_state.page == 2:
    st.title("🔍 Konfirmasi Kerangka Modul")
    d = st.session_state.data
    
    with st.container():
        st.markdown("---")
        c1, c2 = st.columns(2)
        with c1:
            st.write(f"**Nama Guru:** {d['nama']}")
            st.write(f"**Mata Pelajaran:** {d['mapel']}")
            st.write(f"**Topik:** {d['topik']}")
            st.write(f"**Dimensi DPL:** {', '.join(d['dimensi_dpl'])}")
        with c2:
            st.write(f"**Fase / Kelas:** {d['fase']} / {d['kelas']}")
            st.write(f"**Model:** {d['model']}")
            st.write(f"**Alokasi:** {d['jp']} ({d['pertemuan']} Pertemuan)")
        
        if d['kondisi_khusus']:
            st.info(f"**Kondisi Tambahan:** {d['kondisi_khusus']}")
        st.markdown("---")

    col_nav1, col_nav2 = st.columns(2)
    with col_nav1:
        if st.button("⬅️ Kembali & Perbaiki"): go_to_page(1)
    with col_nav2:
        if st.button("🚀 Generate Modul Ajar Sekarang"): go_to_page(3)

# --- HALAMAN 3: GENERATE ---
elif st.session_state.page == 3:
    st.title("✨ Modul Ajar Deep Learning")
    d = st.session_state.data
    
    try:
        genai.configure(api_key=d['api_key'])
        
        # --- SMART MODEL PICKER UNTUK MENGATASI ERROR 404 ---
        with st.spinner("Menyelaraskan model AI dengan standar PPA 2026..."):
            available_models = [m.name for m in genai.list_models() if 'generateContent' in m.supported_generation_methods]
            
            if "models/gemini-1.5-flash-latest" in available_models:
                selected_model = "models/gemini-1.5-flash-latest"
            elif "models/gemini-1.5-flash" in available_models:
                selected_model = "models/gemini-1.5-flash"
            elif "models/gemini-pro" in available_models:
                selected_model = "models/gemini-pro"
            else:
                selected_model = available_models[0]
            
            model = genai.GenerativeModel(selected_model)

        # PROMPT ENGINEERING STANDAR 2026 DENGAN INTEGRASI KONDISI KHUSUS
        prompt = f"""
        Bertindaklah sebagai Guru Ahli Kurikulum 2026. Buatlah **Modul Ajar** lengkap dengan pendekatan **Deep Learning** (Mindful, Meaningful, Joyful).
        
        IDENTITAS: 
        Nama: {d['nama']}, Unit: {d['unit']}, Mapel: {d['mapel']}, Fase/Kelas: {d['fase']}/{d['kelas']}, Semester: {d['semester']}, Alokasi: {d['jp']}, Topik: {d['topik']}.

	INSTRUKSI KHUSUS PEMBAGIAN ALOKASI WAKTU (WAJIB DIPATUHI): jika Fase {d['fase']} = A atau B atau C maka setiap 1 jam pelajaran (JP) = 35 Menit, jika Fase {d['fase']} = D maka setiap 1 jam pelajaran (JP) = 40 Menit, jika Fase {d['fase']} = E atau F maka setiap 1 jam pelajaran (JP) = 45 Menit.

        INSTRUKSI KHUSUS DARI GURU (WAJIB DIINTEGRASIKAN):
        {d['kondisi_khusus'] if d['kondisi_khusus'] else "Tidak ada instruksi tambahan."}

        STRUKTUR MODUL:
        A. CAPAIAN PEMBELAJARAN (CP): Jabarkan elemen dan rumusan CP sesuai topik {d['topik']} yang mengacu pada BSKAP Nomor 046/H/KR/2025 tentang Capaian Pembelajaran (CP) terbaru untuk PAUD, Pendidikan Dasar, dan Menengah (SD, SMP, SMA/SMK) pada Kurikulum Merdeka.
        B. DIMENSI PROFIL LULUSAN (DPL): Integrasikan dimensi {', '.join(d['dimensi_dpl'])} secara eksplisit dalam aktivitas.
	C. CAKUPAN MATERI: rumusan ruang lingkup materi apa saja yang akan dilaksanakan dalam pembelajaran sesuai dengan topik {d['topik']}.        
D. DESAIN PEMBELAJARAN : Terdiri dari 
1. TUJUAN PEMBELAJARAN: Tuliskan  rumusan  tujuan  pembelajaran  apa  yang  akan  dicapai  dalam  pembelajaran  yang mencakup kompetensi dan konten pada ruang lingkup materi dengan menggunakan kata kerja operasional yang relevan, Selan itu dalam merumuskan tujuan pembejaran harus mengandung ABCD, yaitu Audien, Behavior, Condition, dan Degree) Fokus pada kedalaman pemahaman (Deep Learning).
2. PRAKTIK PAEDAGOGIS :Tuliskan model {d['model']} yang dipilih untuk mencapai tujuan pembelajaran dan tuliskan sintaksnya.
3.  KEMITRAAN PEMBELAJARAN (OPSIONAL) :Tuliskan kegiatan kemitraan atau kolaborasi dalam dan/atau ruang lingkup sekolah, seperti: kemitraan antar guru, lintas mata pelajaran, antar murid antar kelas, antar guru lintas sekolah, orang tua, komunitas, tokoh masyarakat, dunia usaha dan dunia industri kerja, institusi, atau mitra profesional.
4. LINGKUNGAN PEMBELAJARAN : Tuliskan lingkungan pembelajaran yang diinginkan dalam pembelajaran dalam budaya belajar, ruang fisik dan/atau ruang virtual agar tecipta iklim belajar yang aman, nyaman, dan saling memuliakan, contoh : memberikan kepada siswa untuk menyampaikan pendapatnya dalam ruang kelas dan dan forum diskusi pada platform daring (ruang virtual bersifat opsional).
5. PEMANFAATAN DIGITAL (OPSIONAL):Tuliskan pemanfaatan digital untuk menciptakan pembelajaran yang inteaktif, kolaboratif dan kontekstual, contoh : video pembelajaran, platform pembelajaran, perpustakaan digital, forum diskusi daing, aplikasi penilaian, dan sebagainya.

        E. PEMAHAMAN BERMAKNA & PERTANYAAN PEMANTIK: 3 Pertanyaan HOTS.
        F. LANGKAH-LANGKAH PEMBELAJARAN (Sintaks {d['model']} buat dalam {d['pertemuan']} pertemuan):
           Wajib mencakup 3 Kategori Deep Learning:
           1. MEMAHAMI (Berkesadaran & Bermakna)
           2. MENGAPLIKASI (Berkesadaran, Bermakna, Menyenangkan)
           3. MEREFLEKSI (Berkesadaran & Bermakna)
           Rincian: a. Pendahuluan, b. Inti (Sintaks {d['model']}), c. Penutup.
        G. ASESMEN: 
        WAJIB: Sajikan bagian asesmen dalam TABEL TERPISAH dengan ketentuan sebagai berikut:
        1. INSTRUMEN ASESMEN : Tuliskan instrumen asesment yang akan dipergunakan selama proses pembelajaran berlangsung dai awal sampai akhir Sajikan dalam tabel.
        2. TEKHNIK ASESMEN: Tuliskan tehnik asesment yang akan dipergunakan selama proses pembelajaran berlangsung dai awal  sampai  akhir,  apakah  menggunakan  tehnik  tes,  yaitu  :  tes  tulis,  tes  lisan,  atau  tes perbuatan dan non tes, yaitu : penilaian sejawat, penilaian diri, penilaian produk, observasi, portofolio, penilaian berbasis kelas, penilaian kinerja, skala sikap, wawancara, atau sosiometri, sajikan dalam tabel.
        
	
	H. MEDIA, ALAT, DAN SUMBER BELAJAR : 
	1. MEDIA DAN ALAT PEMBELAJARAN : Tuliskan media dan alat pembelajaran yang akan dipergunakan pada saat pembelajaran berlangsung untuk membantu dan/atau mempermudah pemahaman murid dalam menerima materi pembelajaran.
	2. SUMBER BELAJAR : Tuliskan referensi baik berupa buku, jurnal, kamus, surat kabar, majalah, website, dan/atau yang lainnya yang akan  dipakai  selama proses  pembelajaran  dalam mendukung  ketecapaian kompetensi seperti yang telah dirumuskan dalam tujuan pembelajaran di atas. Contoh penulisan referensi berupa buku dalam sumber belajar, yaitu : Haris, Mohamad, 2020, Mudah Belajar Matematika, hal. 27-32, edisi kedua, cetakan kesatu, Surabaya, Pelita Bangsa.

	I. LEMBAR KERJA PESERTA DIDIK (LKPD): 
	   1. LKPD : buatkan LKPD sesuai dengan jumlah pertemuan dan Buat instruksi tugas yang jelas dan mendalam setiap LKPD nya disertai RUBRIK PENILAIAN LKPD nya.

        J. LAMPIRAN: Ringkasan Materi Mendalam, dan Glosarium.

        Gunakan format Markdown yang rapi dengan tabel untuk bagian Asesmen.
        """

        with st.status("AI sedang merancang pembelajaran mendalam...", expanded=True) as status:
            response = model.generate_content(prompt)
            status.update(label="Modul Ajar Berhasil Disusun!", state="complete", expanded=False)

        st.markdown(response.text)
        
        # EXPORT KE DOCX
        doc = create_formatted_docx(response.text, f"Modul Ajar - {d['topik']}")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        st.divider()
        c_dl, c_new = st.columns([3, 1])
        with c_dl:
            st.download_button(
                label="📥 Download Modul Ajar (.docx)",
                data=buffer,
                file_name=f"Modul_Ajar_{d['topik'].replace(' ', '_')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        with c_new:
            if st.button("➕ Buat Baru"):
                go_to_page(1)

    except Exception as e:
        st.error(f"Terjadi kendala sistem: {e}")
        if st.button("Coba Hubungkan Kembali"):
            go_to_page(1)
